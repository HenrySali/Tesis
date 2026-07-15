"""
Genera un archivo Word (.docx) a partir de todos los HTML de la tesis.
Usa python-docx + BeautifulSoup para parsear HTML y crear el documento.

Ejecutar: python generar-word.py
Resultado: tesis-smarttemp-v23.docx
"""

import os
import re
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from html.parser import HTMLParser

# Orden de archivos (mismo que index.html)
ARCHIVOS = [
    'portada.html',
    'resumen.html',
    'capitulo-1.html',
    'capitulo-2.html',
    'capitulo-3.html',
    'capitulo-4.html',
    'capitulo-5.html',
    'capitulo-6.html',
    'capitulo-7.html',
    'referencias.html',
    'anexo-a.html',
    'anexo-b.html',
    'anexo-c.html',
    'anexo-d.html',
    'anexo-e.html',
    'anexo-f.html',
    'anexo-g.html',
    'anexo-h.html',
    'anexo-i.html',
    'anexo-k.html',
    'anexo-l.html',
]

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT = os.path.join(BASE_DIR, 'tesis-smarttemp-v26-APA.docx')


class HTMLToDocx(HTMLParser):
    """Parser simple de HTML a python-docx"""
    
    def __init__(self, doc):
        super().__init__()
        self.doc = doc
        self.current_paragraph = None
        self.in_table = False
        self.table_data = []
        self.current_row = []
        self.current_cell = ''
        self.bold = False
        self.italic = False
        self.in_code = False
        self.in_li = False
        self.list_counter = 0
        self.skip = False
        self.tag_stack = []
    
    def handle_starttag(self, tag, attrs):
        self.tag_stack.append(tag)
        attrs_dict = dict(attrs)
        
        if tag == 'h1':
            self.current_paragraph = self.doc.add_heading('', level=1)
        elif tag == 'h2':
            self.current_paragraph = self.doc.add_heading('', level=2)
        elif tag == 'h3':
            self.current_paragraph = self.doc.add_heading('', level=3)
        elif tag == 'p':
            if not self.in_table:
                self.current_paragraph = self.doc.add_paragraph()
        elif tag == 'strong' or tag == 'b':
            self.bold = True
        elif tag == 'em' or tag == 'i':
            self.italic = True
        elif tag == 'code':
            self.in_code = True
        elif tag == 'pre':
            self.current_paragraph = self.doc.add_paragraph()
            self.current_paragraph.style = self.doc.styles['No Spacing']
            self.current_paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            self.current_paragraph.paragraph_format.first_line_indent = Cm(0)
            self.in_code = True
        elif tag == 'table':
            self.in_table = True
            self.table_data = []
        elif tag == 'tr':
            self.current_row = []
        elif tag == 'td' or tag == 'th':
            self.current_cell = ''
        elif tag == 'ul' or tag == 'ol':
            self.list_counter = 0
        elif tag == 'li':
            self.list_counter += 1
            self.current_paragraph = self.doc.add_paragraph(style='List Bullet')
            self.in_li = True
        elif tag == 'img':
            src = attrs_dict.get('src', '')
            alt = attrs_dict.get('alt', '')
            img_path = os.path.join(BASE_DIR, src)
            if os.path.exists(img_path):
                try:
                    self.doc.add_picture(img_path, width=Inches(5))
                    last_paragraph = self.doc.paragraphs[-1]
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as e:
                    p = self.doc.add_paragraph(f'[Imagen: {src}]')
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p = self.doc.add_paragraph(f'[Imagen no encontrada: {src}]')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif tag == 'br':
            if self.current_paragraph:
                self.current_paragraph.add_run('\n')
        elif tag in ('div', 'section', 'span', 'thead', 'tbody'):
            pass  # ignorar
        elif tag == 'script' or tag == 'style':
            self.skip = True
    
    def handle_endtag(self, tag):
        if self.tag_stack and self.tag_stack[-1] == tag:
            self.tag_stack.pop()
        
        if tag == 'script' or tag == 'style':
            self.skip = False
        elif tag == 'strong' or tag == 'b':
            self.bold = False
        elif tag == 'em' or tag == 'i':
            self.italic = False
        elif tag == 'code' and not any(t == 'pre' for t in self.tag_stack):
            self.in_code = False
        elif tag == 'pre':
            self.in_code = False
        elif tag == 'td' or tag == 'th':
            self.current_row.append(self.current_cell.strip())
        elif tag == 'tr':
            if self.current_row:
                self.table_data.append(self.current_row)
        elif tag == 'table':
            self.in_table = False
            if self.table_data:
                self._add_table()
        elif tag == 'li':
            self.in_li = False
    
    def handle_data(self, data):
        if self.skip:
            return
        
        text = data
        
        if self.in_table:
            self.current_cell += text
            return
        
        # En bloques de código, preservar texto tal cual (con saltos de línea)
        if self.in_code:
            if self.current_paragraph is None:
                self.current_paragraph = self.doc.add_paragraph()
                self.current_paragraph.style = self.doc.styles['No Spacing']
                self.current_paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                self.current_paragraph.paragraph_format.first_line_indent = Cm(0)
            # Agregar cada línea como run separado con salto de línea
            lines = text.split('\n')
            for i, line in enumerate(lines):
                if i > 0:
                    self.current_paragraph.add_run('\n')
                if line:
                    run = self.current_paragraph.add_run(line)
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
            return
        
        if not text.strip():
            return
            
        if self.current_paragraph is None:
            self.current_paragraph = self.doc.add_paragraph()
        
        run = self.current_paragraph.add_run(text)
        run.bold = self.bold
        run.italic = self.italic
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    
    def _add_table(self):
        """Agrega una tabla al documento"""
        if not self.table_data:
            return
        
        num_cols = max(len(row) for row in self.table_data)
        table = self.doc.add_table(rows=0, cols=num_cols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        for i, row_data in enumerate(self.table_data):
            row = table.add_row()
            for j, cell_text in enumerate(row_data):
                if j < num_cols:
                    row.cells[j].text = cell_text
                    # Primera fila en negrita (encabezado)
                    if i == 0:
                        for paragraph in row.cells[j].paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
        
        # Espacio después de tabla
        self.doc.add_paragraph()
        self.table_data = []


def generar_word():
    print("Generando Word de la tesis completa...")
    doc = Document()
    
    # Configurar márgenes APA 7: 2.54 cm (1 pulgada) todos los lados
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
    
    # Configurar estilo Normal: Times New Roman 12pt, interlineado doble
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    pf = style.paragraph_format
    pf.line_spacing = 2.0  # Interlineado doble
    pf.space_after = Pt(0)
    pf.first_line_indent = Cm(1.27)  # Sangría primera línea APA
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Configurar estilo No Spacing para código: Courier New 9pt, sin sangría, izquierda
    no_spacing = doc.styles['No Spacing']
    no_spacing.font.name = 'Courier New'
    no_spacing.font.size = Pt(9)
    no_spacing.paragraph_format.line_spacing = 1.0
    no_spacing.paragraph_format.space_after = Pt(0)
    no_spacing.paragraph_format.space_before = Pt(0)
    no_spacing.paragraph_format.first_line_indent = Cm(0)
    no_spacing.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Configurar estilos de encabezados
    for level in range(1, 4):
        h_style = doc.styles[f'Heading {level}']
        h_style.font.name = 'Times New Roman'
        h_style.font.size = Pt(12)
        h_style.font.bold = True
        h_style.paragraph_format.line_spacing = 2.0
        h_style.paragraph_format.space_before = Pt(12)
        h_style.paragraph_format.space_after = Pt(0)
        h_style.paragraph_format.first_line_indent = Cm(0)
        if level == 1:
            h_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            h_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if level == 3:
            h_style.font.italic = True
    
    archivos_procesados = 0
    
    for archivo in ARCHIVOS:
        filepath = os.path.join(BASE_DIR, archivo)
        if not os.path.exists(filepath):
            print(f"  ⚠️  No encontrado: {archivo}")
            continue
        
        print(f"  Procesando: {archivo}")
        
        with open(filepath, 'r', encoding='utf-8') as f:
            html_content = f.read()
        
        # Salto de página entre secciones
        if archivos_procesados > 0:
            doc.add_page_break()
        
        parser = HTMLToDocx(doc)
        parser.feed(html_content)
        archivos_procesados += 1
    
    # Guardar
    doc.save(OUTPUT)
    print(f"\n✅ Documento generado: {OUTPUT}")
    print(f"   Archivos procesados: {archivos_procesados}/{len(ARCHIVOS)}")


if __name__ == '__main__':
    generar_word()
